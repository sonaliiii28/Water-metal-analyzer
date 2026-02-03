import streamlit as st
import pandas as pd
import numpy as np
from sklearn.decomposition import PCA
import matplotlib.pyplot as plt
from openai import OpenAI
import os
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import io

# ----------------------------
# Page Setup
# ----------------------------
st.set_page_config(page_title="WaterMetal Analyzer", layout="wide")
st.title("🌊 WaterMetal Analyzer – Heavy Metal Risk Assessment")

# ----------------------------
# OpenAI Client (optional)
# ----------------------------
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ----------------------------
# File Upload
# ----------------------------
uploaded_file = st.file_uploader("Upload Excel or CSV file", type=["xlsx","csv"])

if uploaded_file:

    # Load Data
    if uploaded_file.name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)

    st.subheader("📄 Uploaded Data")
    st.dataframe(df)

    # Metals
    metals = ["Fe","Mn","Cr","Cu","Ni","Co","Pb","Zn"]

    # PERI Parameters
    background = {"Fe":35000,"Mn":600,"Cr":90,"Cu":45,"Ni":50,"Co":19,"Pb":20,"Zn":95}
    toxicity = {"Fe":1,"Mn":1,"Cr":5,"Cu":5,"Ni":5,"Co":5,"Pb":10,"Zn":1}

    # PERI Calculation
    peri = []
    for i,row in df.iterrows():
        p = 0
        for m in metals:
            p += toxicity[m] * (row[m] / background[m])
        peri.append(p)

    df["PERI"] = peri

    st.subheader("☠️ Ecological Risk (PERI)")
    st.dataframe(df[["S.No","PERI"]])

    # Hotspots
    hotspots = df.sort_values("PERI", ascending=False).head(5)
    st.subheader("🔥 Top 5 High-Risk Stations")
    st.table(hotspots[["S.No","PERI"]])

    # PCA
    X = (df[metals] - df[metals].mean()) / df[metals].std()
    pca = PCA(n_components=2)
    pcs = pca.fit_transform(X)

    df["PC1"] = pcs[:,0]
    df["PC2"] = pcs[:,1]

    st.subheader("🧠 PCA Pollution Pattern")
    fig, ax = plt.subplots()
    ax.scatter(df["PC1"], df["PC2"])
    ax.set_xlabel("PC1")
    ax.set_ylabel("PC2")
    st.pyplot(fig)

    # Metal Risk
    metal_risk = {}
    for m in metals:
        metal_risk[m] = (toxicity[m] * df[m] / background[m]).sum()

    metal_df = pd.DataFrame({
        "Metal": metal_risk.keys(),
        "Total Risk": metal_risk.values()
    })
    metal_df["Percent"] = 100 * metal_df["Total Risk"] / metal_df["Total Risk"].sum()

    st.subheader("🔬 Metal-wise Risk Contribution")
    st.dataframe(metal_df)

    # ----------------------------
    # AI Chatbot (Optional)
    # ----------------------------
    st.subheader("🤖 Ask the AI about your results")

    user_question = st.text_input("Ask something about the pollution:")

    if user_question:
        try:
            context = f"""
            PERI values:
            {df[['S.No','PERI']].to_string()}

            Metal risk:
            {metal_df.to_string()}

            Hotspots:
            {hotspots[['S.No','PERI']].to_string()}
            """

            prompt = f"""
            You are an environmental scientist.
            Use the data below to answer the user's question.

            {context}

            Question: {user_question}
            """

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role":"system","content":"You are an environmental scientist."},
                    {"role":"user","content":prompt}
                ]
            )

            st.success(response.choices[0].message.content)

        except:
            st.warning("AI not available (no API credits). Core analysis is working.")

    # ----------------------------
    # DOWNLOAD REPORTS
    # ----------------------------
    st.subheader("📥 Download Analysis Reports")

    # WORD
    if st.button("Generate Word Report"):
        doc = Document()
        doc.add_heading("WaterMetal Analyzer – Heavy Metal Risk Report", 0)

        doc.add_paragraph("Top 5 High-Risk Stations:")
        for i, row in hotspots.iterrows():
            doc.add_paragraph(f"Station {int(row['S.No'])} : PERI = {row['PERI']:.2f}")

        doc.add_paragraph("\nMetal Risk Contribution:")
        for i, row in metal_df.iterrows():
            doc.add_paragraph(f"{row['Metal']} : {row['Percent']:.2f}%")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="Download Word (.docx)",
            data=buffer,
            file_name="WaterMetal_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # PDF
    if st.button("Generate PDF Report"):
        buffer = io.BytesIO()
        styles = getSampleStyleSheet()
        pdf = SimpleDocTemplate(buffer)

        story = []
        story.append(Paragraph("WaterMetal Analyzer – Heavy Metal Risk Report", styles["Title"]))
        story.append(Paragraph("Top 5 High-Risk Stations:", styles["Heading2"]))

        for i, row in hotspots.iterrows():
            story.append(Paragraph(f"Station {int(row['S.No'])} : PERI = {row['PERI']:.2f}", styles["Normal"]))

        story.append(Paragraph("Metal Risk Contribution:", styles["Heading2"]))
        for i, row in metal_df.iterrows():
            story.append(Paragraph(f"{row['Metal']} : {row['Percent']:.2f}%", styles["Normal"]))

        pdf.build(story)
        buffer.seek(0)

        st.download_button(
            label="Download PDF",
            data=buffer,
            file_name="WaterMetal_Report.pdf",
            mime="application/pdf"
        )

    # EXCEL
    if st.button("Generate Excel Summary"):
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Full Data", index=False)
            metal_df.to_excel(writer, sheet_name="Metal Risk", index=False)
            hotspots.to_excel(writer, sheet_name="Hotspots", index=False)

        excel_buffer.seek(0)

        st.download_button(
            label="Download Excel",
            data=excel_buffer,
            file_name="WaterMetal_Report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
